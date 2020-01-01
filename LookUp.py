import os
import docx



look = input("Enter the keyword: ")
look=str(look)
#print(os.listdir(r"C:\Users\Richard\Desktop\LookupTest"))
Results=[]
def searchThroughFileDocx(filename):
    doc=docx.Document(filename)
    FullText=""
    for z in range(len(doc.paragraphs)):
        FullText+= " " + doc.paragraphs[z].text
    if look in FullText:
          Results.append(filename)

def searchThroughFile(filename):    
    with open(filename,encoding="Latin-1") as l:
        if look in l.read():
          Results.append(l)
        else:
            print(f"{l} is not one of them")


path=os.path.join(os.path.join(os.environ['USERPROFILE']), 'Documents') #folder holding the documents
#os.chdir(path)

for i in os.listdir(path):
    New=os.path.join(path,i)
    if os.path.isdir(New):
        try:
            for y in os.listdir(New):
                print(f"file is {y}")
                if "docx" in y:
                    searchThroughFileDocx(y)
                else:
                    searchThroughFile(y)
        except Exception as e:
            print(f"File {New} has an {e} exception ")
    print(f"file containing word is {New}")
    #if os.path.isfile(New):
        #print(f"file is def a file{New}")
    try:
        if "docx" in i:
            searchThroughFileDocx(os.path.join(path,New))
        else:
            searchThroughFile(os.path.join(path,New))
    except Exception as e:
        print(f"File {New} has an {e} exception ")
print("\n\n\n")

for x,y in enumerate(Results):
    print(str(x) + ", " + y)
